"""
Tool: create_rfp_brief_doc

Create a structured Word document (Bid Intelligence Brief) from the synthesised
RFP evaluation content. The document is saved to a project-specific subfolder
inside the configured OneDrive RFP Projects folder, which OneDrive auto-syncs
to the cloud.

The document contains four visually distinct parts:
  Part 1 — Bid Intelligence Brief (9 sections + bid recommendation)
  Part 2 — Clarification Questions (submit by Q&A deadline)
  Part 3 — Proposal Draft: Case Studies (Section 6.1, ready to paste)
  Part 4 — Proposal Draft: Risk Management (Section 6.3, ready to paste)

Configuration keys expected in hub_config (or .env):
  RFP_OUTPUT_FOLDER — Base path for RFP briefs, e.g.:
                      "C:/Users/{name}/OneDrive - Contoso/RFP Projects"
                      Defaults to ~/OneDrive/RFP Projects if not set.
"""

import logging
import os
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger("hub_se_agent")

SCHEMA = {
    "type": "function",
    "name": "create_rfp_brief_doc",
    "description": (
        "Create a formatted Word document containing the Bid Intelligence Brief "
        "for an RFP. Saves to a project-specific subfolder inside the OneDrive "
        "RFP Projects folder so it is automatically synced to the cloud."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "rfp_id": {
                "type": "string",
                "description": "RFP identifier, e.g. 'RFP-2026-NMS-0042'.",
            },
            "client_name": {
                "type": "string",
                "description": "Client company name, e.g. 'Nexagen Mobility Systems'.",
            },
            "brief_content": {
                "type": "string",
                "description": (
                    "Full Bid Intelligence Brief in markdown format. Should contain "
                    "all 7 sections plus the BID RECOMMENDATION, as synthesised from "
                    "FoundryIQ and Fabric Data Agent outputs."
                ),
            },
            "submission_deadline": {
                "type": "string",
                "description": "Proposal submission deadline, e.g. 'May 16, 2026'.",
            },
        },
        "required": ["rfp_id", "client_name", "brief_content"],
    },
}

# Colour palette
_BRAND_BLUE  = RGBColor(0x1F, 0x49, 0x7D)    # dark navy  — main headings
_ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)    # mid blue   — section banners
_TEAL        = RGBColor(0x00, 0x70, 0x70)    # teal       — clarification questions
_GREEN       = RGBColor(0x37, 0x5C, 0x23)    # dark green — proposal draft headings
_LIGHT_GREY  = "F2F2F2"                       # table header fill
_LIGHT_TEAL  = "E2EFDA"                       # clarification Q background
_LIGHT_GREEN = "E2EFDA"                       # draft section background
_WARN_AMBER  = RGBColor(0xC5, 0x5A, 0x11)    # risk / caution flags

# Sentinel strings the agent embeds in its markdown to mark part boundaries
_SENTINEL_CLARIFICATION = "CLARIFICATION QUESTIONS"
_SENTINEL_DRAFT_A       = "DRAFT A"
_SENTINEL_DRAFT_B       = "DRAFT B"


def _set_run_font(run, size_pt: float, bold=False, italic=False,
                  color: RGBColor | None = None):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = tcBorders.makeelement(qn(f"w:{edge}"), {})
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")


def _add_section_heading(doc: Document, text: str, level: int = 2):
    """Add a styled section heading."""
    para = doc.add_paragraph()
    run = para.add_run(text.upper() if level == 2 else text)
    size = 13 if level == 2 else 11
    color = _BRAND_BLUE if level == 2 else _ACCENT_BLUE
    _set_run_font(run, size, bold=True, color=color)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    # Bottom border on section headings
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    pPr.append(pBdr)
    bottom = pBdr.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    return para


def _add_key_value_row(doc: Document, key: str, value: str):
    """Add a bold-key : value line for metadata."""
    para = doc.add_paragraph()
    run_k = para.add_run(f"{key}: ")
    _set_run_font(run_k, 10, bold=True)
    run_v = para.add_run(value)
    _set_run_font(run_v, 10)
    para.paragraph_format.space_after = Pt(2)


def _render_markdown_body(doc: Document, markdown: str):
    """
    Convert a markdown string into Word paragraphs.
    Handles: # headings, ## headings, **bold**, bullet lists, plain text.
    Does not attempt to render markdown tables — those come from Fabric
    already formatted as prose in the agent response.
    """
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            _add_section_heading(doc, stripped[2:].strip(), level=1)
            continue

        # H2 / H3
        if stripped.startswith("### "):
            _add_section_heading(doc, stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            _add_section_heading(doc, stripped[3:].strip(), level=2)
            continue

        # Bullet
        if stripped.startswith(("- ", "* ", "\u2022 ")):
            para = doc.add_paragraph(style="List Bullet")
            _render_inline(para, stripped[2:].strip())
            para.paragraph_format.space_after = Pt(2)
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            para = doc.add_paragraph(style="List Number")
            _render_inline(para, re.sub(r"^\d+\.\s", "", stripped))
            para.paragraph_format.space_after = Pt(2)
            continue

        # Horizontal rule
        if stripped in ("---", "___", "***"):
            doc.add_paragraph()
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        _render_inline(para, stripped)
        para.paragraph_format.space_after = Pt(3)


def _render_inline(para, text: str):
    """Render inline markdown (bold, italic) into a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*|\*[^*]+?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            _set_run_font(run, 10, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            _set_run_font(run, 10, italic=True)
        else:
            run = para.add_run(part)
            _set_run_font(run, 10)


def _add_part_banner(doc: Document, part_number: str, title: str,
                     subtitle: str, fill_hex: str, text_color: RGBColor):
    """
    Add a full-width shaded banner that visually separates the four document
    parts (Brief / Clarification Questions / Draft A / Draft B).
    """
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill_hex)
    _set_cell_borders(cell)

    # Part label (small caps style)
    p1 = cell.paragraphs[0]
    run = p1.add_run(f"PART {part_number}  \u00b7  {title.upper()}")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = text_color
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)

    # Subtitle
    p2 = cell.add_paragraph()
    run2 = p2.add_run(subtitle)
    run2.font.name = "Calibri"
    run2.font.size = Pt(9)
    run2.italic = True
    run2.font.color.rgb = text_color
    p2.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()


def _split_brief_content(markdown: str) -> tuple[str, str, str, str]:
    """
    Split the agent's markdown output into four parts using sentinel headings.
    Returns (brief_text, clarification_text, draft_a_text, draft_b_text).
    Parts that are not present return empty strings.
    """
    # Normalise line endings
    text = markdown.replace("\r\n", "\n")

    def _extract_between(text: str, start_sentinel: str, end_sentinel: str | None) -> tuple[str, str]:
        """Extract content between two sentinels; return (before, extracted)."""
        pattern = re.compile(
            rf"(?:^#+\s*{re.escape(start_sentinel)}[^\n]*$)",
            re.IGNORECASE | re.MULTILINE,
        )
        m = pattern.search(text)
        if not m:
            return text, ""
        before = text[: m.start()].rstrip()
        after = text[m.start():]
        if end_sentinel:
            end_pattern = re.compile(
                rf"(?:^#+\s*{re.escape(end_sentinel)}[^\n]*$)",
                re.IGNORECASE | re.MULTILINE,
            )
            m2 = end_pattern.search(after, m.end() - m.start())
            if m2:
                extracted = after[: m2.start()]
                remainder = after[m2.start():]
                return before + "\n" + remainder, extracted
        return before, after

    # Peel off Draft B first, then Draft A, then Clarification Questions
    text, draft_b = _extract_between(text, _SENTINEL_DRAFT_B, None)
    text, draft_a = _extract_between(text, _SENTINEL_DRAFT_A, None)
    brief, clarification = _extract_between(text, _SENTINEL_CLARIFICATION, None)

    return brief.strip(), clarification.strip(), draft_a.strip(), draft_b.strip()


def _add_cover_metadata(doc: Document, rfp_id: str, client_name: str,
                         submission_deadline: str):
    """Add a styled cover block at the top of the document."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("BID INTELLIGENCE BRIEF")
    _set_run_font(run, 22, bold=True, color=_BRAND_BLUE)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(client_name)
    _set_run_font(run, 14, bold=True, color=_ACCENT_BLUE)
    subtitle.paragraph_format.space_after = Pt(8)

    # Metadata table (2 cols)
    meta_rows = [
        ("RFP Reference", rfp_id),
        ("Client", client_name),
        ("Prepared by", "Contoso Engineering \u2014 Project Intelligence Agent"),
        ("Date prepared", date.today().strftime("%B %d, %Y")),
        ("Submission deadline", submission_deadline or "See RFP document"),
        ("Classification", "CONFIDENTIAL \u2014 Internal Use Only"),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    col_widths = [Cm(5), Cm(11)]
    for i, (key, value) in enumerate(meta_rows):
        row = table.rows[i]
        # Key cell
        key_cell = row.cells[0]
        key_cell.text = ""
        run = key_cell.paragraphs[0].add_run(key)
        _set_run_font(run, 9, bold=True)
        _set_cell_shading(key_cell, _LIGHT_GREY)
        _set_cell_borders(key_cell)
        # Value cell
        val_cell = row.cells[1]
        val_cell.text = ""
        run = val_cell.paragraphs[0].add_run(value)
        _set_run_font(run, 9)
        _set_cell_borders(val_cell)
        for j, w in enumerate(col_widths):
            row.cells[j].width = w

    doc.add_paragraph()  # spacing after cover block


def _get_output_folder(client_name: str, rfp_id: str) -> Path:
    """Resolve the output folder from hub_config or default to OneDrive."""
    base = ""
    try:
        import hub_config
        cfg = hub_config.load()
        base = cfg.get("RFP_OUTPUT_FOLDER", "")
    except Exception:
        pass
    if not base:
        base = os.environ.get("RFP_OUTPUT_FOLDER", "")
    if not base:
        # Default: ~/OneDrive/RFP Projects  (works for personal OneDrive)
        onedrive = Path.home() / "OneDrive"
        if not onedrive.exists():
            # Try OneDrive - CompanyName variant
            candidates = list(Path.home().glob("OneDrive*"))
            onedrive = candidates[0] if candidates else Path.home() / "Documents"
        base = str(onedrive / "RFP Projects")

    # Sanitise client name and rfp_id for use as folder name
    safe_client = re.sub(r'[\\/*?:"<>|]', "", client_name).strip()
    safe_rfp = re.sub(r'[\\/*?:"<>|]', "", rfp_id).strip()
    folder_name = f"{safe_client} \u2014 {safe_rfp}"
    output_path = Path(base) / folder_name
    output_path.mkdir(parents=True, exist_ok=True)
    return output_path


def handle(arguments: dict, *, on_progress=None, **kwargs) -> str:
    """Create the Bid Intelligence Brief Word document."""
    rfp_id = arguments["rfp_id"]
    client_name = arguments["client_name"]
    brief_content = arguments["brief_content"]
    submission_deadline = arguments.get("submission_deadline", "")

    if on_progress:
        on_progress("tool", f"Creating RFP Brief document for {client_name} / {rfp_id}")

    # Build filename and output path
    safe_client = re.sub(r'[\\/*?:"<>|]', "", client_name).strip().replace(" ", "-")
    safe_rfp = re.sub(r'[\\/*?:"<>|]', "", rfp_id).strip()
    filename = f"RFP-Brief-{safe_client}-{safe_rfp}.docx"

    output_folder = _get_output_folder(client_name, rfp_id)
    file_path = output_folder / filename

    try:
        doc = Document()

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)

        # Split the agent output into its four parts
        brief_text, clarification_text, draft_a_text, draft_b_text = \
            _split_brief_content(brief_content)

        # ── PART 1: Bid Intelligence Brief ───────────────────────────────────
        _add_part_banner(
            doc,
            part_number="1",
            title="Bid Intelligence Brief",
            subtitle=(
                "Internal analysis \u2014 do not distribute externally. "
                "Synthesised from FoundryIQ case study narratives and Fabric OneLake project data."
            ),
            fill_hex="1F497D",
            text_color=RGBColor(0xFF, 0xFF, 0xFF),
        )
        _add_cover_metadata(doc, rfp_id, client_name, submission_deadline)
        _render_markdown_body(doc, brief_text or brief_content)

        # ── PART 2: Clarification Questions ─────────────────────────────────
        if clarification_text:
            doc.add_page_break()
            _add_part_banner(
                doc,
                part_number="2",
                title="Clarification Questions",
                subtitle=(
                    "Submit these questions to the client by the Q&A deadline. "
                    "Each question is grounded in a specific gap or past project risk."
                ),
                fill_hex="006060",
                text_color=RGBColor(0xFF, 0xFF, 0xFF),
            )
            _render_markdown_body(doc, clarification_text)

        # ── PART 3: Proposal Draft — Case Studies ────────────────────────────
        if draft_a_text:
            doc.add_page_break()
            _add_part_banner(
                doc,
                part_number="3",
                title="Proposal Draft \u2014 Case Studies (Section 6.1)",
                subtitle=(
                    "Ready-to-paste proposal language. Review and personalise before submission. "
                    "Narratives from FoundryIQ; KPI figures from Fabric OneLake."
                ),
                fill_hex="375C23",
                text_color=RGBColor(0xFF, 0xFF, 0xFF),
            )
            _render_markdown_body(doc, draft_a_text)

        # ── PART 4: Proposal Draft — Risk Management ─────────────────────────
        if draft_b_text:
            doc.add_page_break()
            _add_part_banner(
                doc,
                part_number="4",
                title="Proposal Draft \u2014 Risk Management (Section 6.3)",
                subtitle=(
                    "Evidence-based risk examples in proposal language. "
                    "Directly addresses the client's stated requirement for specific, scored examples."
                ),
                fill_hex="375C23",
                text_color=RGBColor(0xFF, 0xFF, 0xFF),
            )
            _render_markdown_body(doc, draft_b_text)

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        run = footer_para.add_run(
            "Prepared by the Contoso Engineering RFP Evaluation Agent \u00b7 "
            f"{date.today().strftime('%B %d, %Y')} \u00b7 "
            "Sources: FoundryIQ (case study PDFs) + Fabric OneLake (structured project data) + "
            "WorkIQ (M365 emails, calendar, SharePoint). "
            "Internal use only \u2014 review before external distribution."
        )
        _set_run_font(run, 8, italic=True, color=RGBColor(0x7F, 0x7F, 0x7F))

        doc.save(str(file_path))
        logger.info("[RFPBrief] Document saved: %s", file_path)

        # Auto-open on Windows
        if sys.platform == "win32":
            try:
                os.startfile(str(file_path))
            except Exception as e:
                logger.warning("[RFPBrief] Could not auto-open: %s", e)

        if on_progress:
            on_progress("tool", f"Document saved: {filename}")

        return (
            f"RFP Brief document created successfully.\n"
            f"Filename : {filename}\n"
            f"Saved to : {file_path}\n"
            f"OneDrive will sync this folder automatically."
        )

    except Exception as e:
        logger.error("[RFPBrief] Failed: %s", e, exc_info=True)
        return f"Error creating RFP Brief document: {e}"
